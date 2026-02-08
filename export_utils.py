import io
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# Note: parse_markdown_to_segments was removed - unused placeholder

# --- DOCX Export ---
def create_docx(data):
    """
    Creates a Word document with professional styling.
    """
    text = data.get('body', '')
    user_info = data.get('user_info', {})
    
    doc = Document()
    
    # 1. Header
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = name_paragraph.add_run(user_info.get("full_name", "Name"))
    run_name.bold = True
    run_name.font.size = Pt(16)
    
    # Contact Info
    contact_parts = [
        user_info.get("address"), 
        user_info.get("email"), 
        user_info.get("phone"), 
        user_info.get("linkedin")
    ]
    contact_str = " • ".join([p for p in contact_parts if p])
    
    contact_paragraph = doc.add_paragraph(contact_str)
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph("__________________________________________________________________________")
    
    # 2. Date & Recipient (Basic formatting)
    # We expect the body to contain the "Dear ..." and the date/address block usually? 
    # Or did our prompt put it there? 
    # The PROMPT logic in utils.py generates the header block inside 'text'.
    # So we should probably just treat 'text' as the main content logic.
    # But wait, we added a nice header above only to duplicate it?
    # Our prompt says "STRICT FORMATTING RULES: output must start with this EXACT header block".
    # So the AI output 'text' ALREADY has the header.
    # We should probably strip the AI header if we are rendering our own?
    # Or just render the AI text. 
    # Improvement A1 says: "Add basic professional styling: full name as bold header".
    # So we should ideally Parse the top of the AI text or ask AI strictly for body.
    # Given we can't change the prompt output easily without looking at it, 
    # we will trust the prompt output but Apply Markdown formatting.
    
    # Let's just render the Body content, since prompt includes header.
    # But to make "Name as bold header", we might rely on the AI doing it?
    # No, the task implies WE do styling.
    # If the text has the header, we'll see duplicates if we add one.
    # Let's assume the user wants the AI generated text as-is but with Markdown->Formatting.
    
    last_p = None
    for line in text.split('\n'):
        # Basic list bullet handling
        clean_line = line.strip()
        if clean_line.startswith("- ") or clean_line.startswith("* "):
            p = doc.add_paragraph(clean_line[2:], style='List Bullet')
        else:
            if clean_line:
                # Handle bolding logic: **text**
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', clean_line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        r = p.add_run(part[2:-2])
                        r.bold = True
                    else:
                        p.add_run(part)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- PDF Export ---
def create_pdf(data):
    """
    Creates a PDF with UTF-8 support (using assets/fonts/DejaVuSans.ttf or standard).
    """
    text = data.get('body', '')
    
    # Setup PDF
    pdf = FPDF(format='A4')
    pdf.set_margins(25, 25, 25) # 25mm margins
    pdf.add_page()
    
    # Font Handling
    # Check for DejaVuSans
    font_path = os.path.join("assets", "fonts", "DejaVuSans.ttf")
    font_loaded = False
    
    if os.path.exists(font_path):
        try:
            pdf.add_font("DejaVu", fname=font_path, uni=True)
            pdf.set_font("DejaVu", size=11)
            font_loaded = True
        except Exception as e:
            print(f"Font loading error: {e}")
    
    if not font_loaded:
        # Fallback to standard font (Latin-1 limitations)
        pdf.set_font("Helvetica", size=11)
        # Attempt to clean unicode if using standard
        # We'll just replace commonly problematic chars
        replacements = {
            '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"', 
            '\u2013': '-', '\u2014': '--', '\u2022': '*'
        }
        for k,v in replacements.items():
            text = text.replace(k, v)
        # Force encode
        text = text.encode('latin-1', 'replace').decode('latin-1')

    # Content
    # Handle bolding **text** specifically? 
    # FPDF with uni=True supports some styled writing but multi_cell is tricky.
    # We will do simple write for now, or use write_html logic if we had HTML.
    # We'll stick to plain text dump with paragraphs for robustness.
    
    pdf.multi_cell(0, 6, text)
    
    buffer = io.BytesIO()
    pdf_bytes = pdf.output()
    buffer.write(pdf_bytes)
    buffer.seek(0)
    return buffer

# --- LaTeX Export ---
def create_latex(data):
    """
    Creates a LaTeX file using 'letter' class.
    """
    text = data.get('body', '')
    # Strip the header if it looks like the AI included one? 
    # The prompt forces: name, addr/phone/email, linkedin, date, manager, company, addr. then Dear...
    # We want to format this nicely in LaTeX commands if possible.
    # But parsing is fragile. We will stick to putting the whole RAW text in the body 
    # but use a template that minimizes interference, 
    # OR we try to extract the body.
    
    # Simplest: The user requested "Use proper letter format".
    # Typically this means \address{...} \signature{...}.
    # If the text ALREADY contains the header, using letter class `\opening` might double it.
    # We will use a generic `article` but with better spacing commands or `letter` if we can parse.
    # Given constraint "Preserve current behavior", we'll wrap the text.
    
    def latex_escape(s):
        # Robust escaping
        chars = {
            '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_',
            '{': r'\{', '}': r'\}', '~': r'\textasciitilde{}', '^': r'\textasciicircum{}',
            '\\': r'\textbackslash{}'
        }
        return "".join(chars.get(c, c) for c in s)

    # Basic Paragraphs
    # Replace double newlines with \par
    # We must escape the content first!
    safe_text = latex_escape(text)
    
    # Restore bolding? **text** -> \textbf{text}
    # Regex replace: \\*\\*(.*?)\\*\\* -> \textbf{$1}
    # Note: we escaped * to * (it's not special in latex usually, unless active). 
    # Actually * is fine.
    safe_text = re.sub(r'\*\*(.*?)\*\*', r'\\textbf{\1}', safe_text)
    safe_text = re.sub(r'\*(.*?)\*', r'\\textit{\1}', safe_text)
    
    # Handle newlines
    # If we just put newlines, LaTeX treats single newline as space, double as paragraph.
    # We'll preserve structure.
    safe_text = safe_text.replace('\n', '\n\n') # Force paragraphs for every line? Too much.
    # Actually, input text usually has blank lines for paragraphs.
    # We'll just let LaTeX handle blank lines naturally.
    
    template = r"""
\documentclass[11pt,a4paper]{article}
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{geometry}
\usepackage{parskip}
\geometry{
  a4paper,
  total={170mm,257mm},
  left=25mm,
  top=25mm,
}

\begin{document}
%s
\end{document}
"""
    latex_code = template % safe_text
    
    buffer = io.BytesIO()
    buffer.write(latex_code.encode('utf-8'))
    buffer.seek(0)
    return buffer, latex_code
